#!/usr/bin/env python3
"""
Renderiza um perfil estruturado (JSON) em curriculo Markdown + DOCX com layout
seguro para ATS: coluna unica, sem tabelas, sem caixas de texto, sem
cabecalho/rodape, fontes padrao.

Uso:
    python render_cv.py perfil.json --idioma pt --saida ./out
    python render_cv.py perfil.json --idioma en --nome-arquivo "Thayse_Oliveira_Data_Scientist_CV"
    python render_cv.py perfil.json --somente-md

Dependencia: python-docx (pip install python-docx). Sem ela, so o Markdown e gerado.
"""

import argparse
import json
import re
import sys
from pathlib import Path

MESES_PT = ["jan", "fev", "mar", "abr", "mai", "jun",
            "jul", "ago", "set", "out", "nov", "dez"]
MESES_EN = ["Jan", "Feb", "Mar", "Apr", "May", "Jun",
            "Jul", "Aug", "Sep", "Oct", "Nov", "Dec"]

ROTULOS = {
    "pt": {
        "resumo": "RESUMO PROFISSIONAL",
        "experiencia": "EXPERIÊNCIA PROFISSIONAL",
        "formacao": "FORMAÇÃO ACADÊMICA",
        "competencias": "COMPETÊNCIAS TÉCNICAS",
        "idiomas": "IDIOMAS",
        "certificacoes": "CERTIFICAÇÕES",
        "projetos": "PROJETOS",
        "publicacoes": "PUBLICAÇÕES",
        "atual": "atual",
    },
    "en": {
        "resumo": "PROFESSIONAL SUMMARY",
        "experiencia": "PROFESSIONAL EXPERIENCE",
        "formacao": "EDUCATION",
        "competencias": "TECHNICAL SKILLS",
        "idiomas": "LANGUAGES",
        "certificacoes": "CERTIFICATIONS",
        "projetos": "PROJECTS",
        "publicacoes": "PUBLICATIONS",
        "atual": "Present",
    },
}


def fmt_data(valor, idioma):
    """AAAA-MM -> 'mar/2023' (pt) ou 'Mar 2023' (en). 'atual' e traduzido."""
    if not valor:
        return ""
    v = str(valor).strip()
    if v.lower() in ("atual", "present", "current", "hoje", "now"):
        return ROTULOS[idioma]["atual"]
    m = re.match(r"^(\d{4})[-/](\d{1,2})$", v)
    if m:
        ano, mes = int(m.group(1)), int(m.group(2))
        if 1 <= mes <= 12:
            if idioma == "pt":
                return f"{MESES_PT[mes-1]}/{ano}"
            return f"{MESES_EN[mes-1]} {ano}"
    return v


def periodo(exp, idioma):
    ini = fmt_data(exp.get("inicio"), idioma)
    fim = fmt_data(exp.get("fim"), idioma)
    if ini and fim:
        return f"{ini} – {fim}"
    return ini or fim or ""


def local(d):
    partes = [d.get("cidade"), d.get("pais")]
    return ", ".join([p for p in partes if p])


def linha_contato(dp):
    itens = [local(dp), dp.get("email"), dp.get("telefone"),
             dp.get("linkedin"), dp.get("github"), dp.get("site")]
    return " · ".join([i for i in itens if i])


# --------------------------------------------------------------------------
# Markdown
# --------------------------------------------------------------------------

def gerar_markdown(perfil, idioma):
    R = ROTULOS[idioma]
    dp = perfil.get("dados_pessoais", {})
    out = [f"# {dp.get('nome', '')}"]
    if dp.get("titulo_profissional"):
        out.append(f"**{dp['titulo_profissional']}**")
    contato = linha_contato(dp)
    if contato:
        out.append(contato)
    extra = [dp.get("autorizacao_trabalho"), dp.get("disponibilidade")]
    extra = [e for e in extra if e]
    if extra:
        out.append(" · ".join(extra))
    out.append("")

    if perfil.get("resumo"):
        out += [f"## {R['resumo']}", "", perfil["resumo"], ""]

    if perfil.get("experiencias"):
        out += [f"## {R['experiencia']}", ""]
        for exp in perfil["experiencias"]:
            out.append(f"### {exp.get('cargo','')} — {exp.get('empresa','')}")
            sub = [local(exp), periodo(exp, idioma)]
            if exp.get("modalidade"):
                sub.append(exp["modalidade"])
            sub = " | ".join([s for s in sub if s])
            if sub:
                out.append(f"*{sub}*")
            out.append("")
            for b in exp.get("bullets", []):
                texto = b.get("texto", "").strip()
                if texto:
                    out.append(f"- {texto}")
            out.append("")

    if perfil.get("formacao"):
        out += [f"## {R['formacao']}", ""]
        for f in perfil["formacao"]:
            titulo = " em ".join([p for p in [f.get("grau"), f.get("curso")] if p]) \
                if idioma == "pt" else " in ".join([p for p in [f.get("grau"), f.get("curso")] if p])
            linha = f"**{titulo}** — {f.get('instituicao','')}"
            det = [local(f), f.get("ano_conclusao")]
            det = " | ".join([d for d in det if d])
            out.append(linha + (f"  \n*{det}*" if det else ""))
            if f.get("detalhe"):
                out.append(f"- {f['detalhe']}")
            out.append("")

    comp = perfil.get("competencias") or {}
    if comp:
        out += [f"## {R['competencias']}", ""]
        for cat, itens in comp.items():
            if itens:
                out.append(f"**{cat}:** " + ", ".join(itens))
        out.append("")

    if perfil.get("idiomas"):
        out += [f"## {R['idiomas']}", ""]
        linhas = []
        for i in perfil["idiomas"]:
            s = f"{i.get('idioma','')} — {i.get('nivel','')}".strip(" —")
            if i.get("certificacao"):
                s += f" ({i['certificacao']})"
            linhas.append(s)
        out.append(" · ".join(linhas))
        out.append("")

    if perfil.get("certificacoes"):
        out += [f"## {R['certificacoes']}", ""]
        for c in perfil["certificacoes"]:
            s = f"- {c.get('nome','')}"
            det = [c.get("emissor"), c.get("ano")]
            det = ", ".join([d for d in det if d])
            if det:
                s += f" — {det}"
            out.append(s)
        out.append("")

    if perfil.get("projetos"):
        out += [f"## {R['projetos']}", ""]
        for p in perfil["projetos"]:
            s = f"- **{p.get('nome','')}** — {p.get('descricao','')}"
            if p.get("tecnologias"):
                s += f" ({', '.join(p['tecnologias'])})"
            if p.get("url"):
                s += f" — {p['url']}"
            out.append(s)
        out.append("")

    if perfil.get("publicacoes"):
        out += [f"## {R['publicacoes']}", ""]
        for p in perfil["publicacoes"]:
            det = ", ".join([d for d in [p.get("veiculo"), p.get("ano")] if d])
            out.append(f"- {p.get('titulo','')}" + (f" — {det}" if det else ""))
        out.append("")

    return "\n".join(out).rstrip() + "\n"


# --------------------------------------------------------------------------
# DOCX (ATS-safe)
# --------------------------------------------------------------------------

def gerar_docx(perfil, idioma, caminho, fonte="Calibri", tamanho=10.5):
    try:
        from docx import Document
        from docx.shared import Pt, Cm, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.ns import qn
    except ImportError:
        print("[aviso] python-docx nao instalado — DOCX nao gerado. "
              "Instale com: pip install python-docx", file=sys.stderr)
        return None

    R = ROTULOS[idioma]
    doc = Document()

    # Margens e fonte base. Sem header/footer: ATS antigo costuma descartar.
    for s in doc.sections:
        s.top_margin = s.bottom_margin = Cm(1.6)
        s.left_margin = s.right_margin = Cm(1.9)

    normal = doc.styles["Normal"]
    normal.font.name = fonte
    normal.font.size = Pt(tamanho)
    normal.element.rPr.rFonts.set(qn("w:eastAsia"), fonte)
    normal.paragraph_format.space_after = Pt(2)
    normal.paragraph_format.line_spacing = 1.05

    def p(texto="", bold=False, italic=False, size=None, align=None,
          space_before=0, space_after=2):
        par = doc.add_paragraph()
        par.paragraph_format.space_before = Pt(space_before)
        par.paragraph_format.space_after = Pt(space_after)
        if align:
            par.alignment = align
        run = par.add_run(texto)
        run.bold = bold
        run.italic = italic
        run.font.name = fonte
        run.font.size = Pt(size or tamanho)
        return par

    def secao(titulo):
        par = doc.add_paragraph()
        par.paragraph_format.space_before = Pt(9)
        par.paragraph_format.space_after = Pt(3)
        run = par.add_run(titulo)
        run.bold = True
        run.font.name = fonte
        run.font.size = Pt(tamanho + 0.5)
        run.font.color.rgb = RGBColor(0x1F, 0x1F, 0x1F)
        # Linha horizontal via borda inferior do paragrafo (nao e tabela nem imagem).
        from docx.oxml import OxmlElement
        pPr = par._p.get_or_add_pPr()
        borders = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "6")
        bottom.set(qn("w:space"), "1")
        bottom.set(qn("w:color"), "999999")
        borders.append(bottom)
        pPr.append(borders)

    def bullet(texto):
        par = doc.add_paragraph(style="List Bullet")
        par.paragraph_format.space_after = Pt(1)
        par.paragraph_format.left_indent = Cm(0.5)
        run = par.add_run(texto)
        run.font.name = fonte
        run.font.size = Pt(tamanho)

    dp = perfil.get("dados_pessoais", {})
    p(dp.get("nome", ""), bold=True, size=tamanho + 6,
      align=WD_ALIGN_PARAGRAPH.CENTER, space_after=1)
    if dp.get("titulo_profissional"):
        p(dp["titulo_profissional"], size=tamanho + 1,
          align=WD_ALIGN_PARAGRAPH.CENTER, space_after=1)
    contato = linha_contato(dp)
    if contato:
        p(contato, size=tamanho - 0.5, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=1)
    extra = [e for e in [dp.get("autorizacao_trabalho"), dp.get("disponibilidade")] if e]
    if extra:
        p(" · ".join(extra), size=tamanho - 0.5, italic=True,
          align=WD_ALIGN_PARAGRAPH.CENTER)

    if perfil.get("resumo"):
        secao(R["resumo"])
        p(perfil["resumo"])

    if perfil.get("experiencias"):
        secao(R["experiencia"])
        for exp in perfil["experiencias"]:
            p(f"{exp.get('cargo','')} — {exp.get('empresa','')}", bold=True,
              space_before=4, space_after=0)
            sub = [local(exp), periodo(exp, idioma)]
            if exp.get("modalidade"):
                sub.append(exp["modalidade"])
            sub = " | ".join([s for s in sub if s])
            if sub:
                p(sub, italic=True, size=tamanho - 0.5, space_after=2)
            for b in exp.get("bullets", []):
                if b.get("texto"):
                    bullet(b["texto"].strip())

    if perfil.get("formacao"):
        secao(R["formacao"])
        for f in perfil["formacao"]:
            juncao = " em " if idioma == "pt" else " in "
            titulo = juncao.join([x for x in [f.get("grau"), f.get("curso")] if x])
            p(f"{titulo} — {f.get('instituicao','')}", bold=True, space_after=0)
            det = " | ".join([d for d in [local(f), f.get("ano_conclusao")] if d])
            if det:
                p(det, italic=True, size=tamanho - 0.5)
            if f.get("detalhe"):
                bullet(f["detalhe"])

    comp = perfil.get("competencias") or {}
    if comp:
        secao(R["competencias"])
        for cat, itens in comp.items():
            if not itens:
                continue
            par = doc.add_paragraph()
            par.paragraph_format.space_after = Pt(1)
            r1 = par.add_run(f"{cat}: ")
            r1.bold = True
            r1.font.name = fonte
            r1.font.size = Pt(tamanho)
            r2 = par.add_run(", ".join(itens))
            r2.font.name = fonte
            r2.font.size = Pt(tamanho)

    if perfil.get("idiomas"):
        secao(R["idiomas"])
        linhas = []
        for i in perfil["idiomas"]:
            s = f"{i.get('idioma','')} — {i.get('nivel','')}".strip(" —")
            if i.get("certificacao"):
                s += f" ({i['certificacao']})"
            linhas.append(s)
        p(" · ".join(linhas))

    if perfil.get("certificacoes"):
        secao(R["certificacoes"])
        for c in perfil["certificacoes"]:
            det = ", ".join([d for d in [c.get("emissor"), c.get("ano")] if d])
            bullet(f"{c.get('nome','')}" + (f" — {det}" if det else ""))

    if perfil.get("projetos"):
        secao(R["projetos"])
        for pr in perfil["projetos"]:
            s = f"{pr.get('nome','')} — {pr.get('descricao','')}"
            if pr.get("tecnologias"):
                s += f" ({', '.join(pr['tecnologias'])})"
            if pr.get("url"):
                s += f" — {pr['url']}"
            bullet(s)

    if perfil.get("publicacoes"):
        secao(R["publicacoes"])
        for pu in perfil["publicacoes"]:
            det = ", ".join([d for d in [pu.get("veiculo"), pu.get("ano")] if d])
            bullet(f"{pu.get('titulo','')}" + (f" — {det}" if det else ""))

    doc.save(caminho)
    return caminho


def slug(texto):
    texto = re.sub(r"[^\w\s-]", "", texto or "curriculo", flags=re.UNICODE)
    return re.sub(r"[\s-]+", "_", texto.strip())


def main():
    ap = argparse.ArgumentParser(description="Renderiza curriculo ATS/Harvard a partir do perfil JSON.")
    ap.add_argument("perfil", help="caminho do perfil.json")
    ap.add_argument("--idioma", default=None, choices=["pt", "en"],
                    help="idioma de saida (default: meta.idioma_saida ou 'pt')")
    ap.add_argument("--saida", default=".", help="diretorio de saida")
    ap.add_argument("--nome-arquivo", default=None,
                    help="nome base do arquivo, sem extensao")
    ap.add_argument("--somente-md", action="store_true", help="nao gerar DOCX")
    ap.add_argument("--fonte", default="Calibri")
    ap.add_argument("--tamanho", type=float, default=10.5)
    args = ap.parse_args()

    perfil = json.loads(Path(args.perfil).read_text(encoding="utf-8"))
    idioma = args.idioma or (perfil.get("meta", {}) or {}).get("idioma_saida") or "pt"

    nome = args.nome_arquivo
    if not nome:
        dp = perfil.get("dados_pessoais", {})
        base = f"{dp.get('nome','curriculo')} {dp.get('titulo_profissional','')} CV"
        nome = slug(base)

    out = Path(args.saida)
    out.mkdir(parents=True, exist_ok=True)

    md = gerar_markdown(perfil, idioma)
    caminho_md = out / f"{nome}.md"
    caminho_md.write_text(md, encoding="utf-8")
    print(f"[ok] Markdown: {caminho_md}")

    if not args.somente_md:
        caminho_docx = out / f"{nome}.docx"
        r = gerar_docx(perfil, idioma, str(caminho_docx), args.fonte, args.tamanho)
        if r:
            print(f"[ok] DOCX: {caminho_docx}")
            print("[dica] Para PDF: soffice --headless --convert-to pdf "
                  f"'{caminho_docx}' (mantem o layout ATS-safe)")

    # Diagnostico rapido de qualidade
    exps = perfil.get("experiencias", [])
    bullets = [b for e in exps for b in e.get("bullets", [])]
    com_metrica = [b for b in bullets if (b.get("metrica") or "").strip()
                   or re.search(r"\d", b.get("texto", ""))]
    if bullets:
        pct = 100 * len(com_metrica) / len(bullets)
        print(f"[diagnostico] {len(bullets)} bullets, {pct:.0f}% com numero/metrica "
              f"(alvo: 60%+)")
    fracos = [b["texto"] for b in bullets
              if re.match(r"^\s*(respons[aá]vel|atuei|participei|auxiliei|ajudei|"
                          r"trabalhei|responsible for|worked on|helped)",
                          b.get("texto", ""), re.IGNORECASE)]
    for f in fracos:
        print(f"[revisar] bullet comeca com verbo fraco: {f[:80]}...")
    lacunas = (perfil.get("meta", {}) or {}).get("lacunas_conhecidas") or []
    for l in lacunas:
        print(f"[lacuna] {l}")


if __name__ == "__main__":
    main()
